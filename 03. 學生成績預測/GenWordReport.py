import pandas as pd
from docx import Document
from docx.shared import Mm
from sklearn.metrics import mean_squared_error
from sklearn.metrics import mean_absolute_error
from sklearn.metrics import median_absolute_error
from sklearn.metrics import r2_score
import matplotlib.pyplot as plt
import os


def replace(putData, train_list):

    for i in train_list:

        if i == 'school':
            putData['school'].replace('GP', 1, inplace=True)
            putData['school'].replace('MS', 2, inplace=True)

        if i == 'sex':
            putData['sex'].replace('F', 1, inplace=True)
            putData['sex'].replace('M', 2, inplace=True)

        if i == 'address':
            putData['address'].replace('U', 1, inplace=True)
            putData['address'].replace('R', 2, inplace=True)


def doc_table(dx, data):

    table = dx.add_table(rows=1, cols=data.shape[1] + 1, style='Light Grid Accent 1')
    row1 = table.rows[0].cells

    for i in range(0, len(data.columns) + 1) :
        if i == 0 : row1[i].text = "index"
        else : row1[i].text = data.columns[i-1]

    for i in list(data.index):
        row = table.add_row().cells
        row[0].text = str(i)
        rowData = [str(x) for x in data.loc[i]]

        for j in range(0, len(data.columns)):
            row[j+1].text = rowData[j]

def genWordReport( choose, train_list, predict, old_data, reportData):

    os.chdir("./File")
    dx = Document()
    section = dx.sections[0]
    section.page_height = Mm(420)
    section.page_width = Mm(297)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(22)

    dx.add_heading('學生成績預測模型訓練報告', 0)  # ----- REPORT INITIALIZATION
    dx.add_paragraph('--這是由學生成績預測系統自動生成的報告')
    dx.add_paragraph('--作者: 林威成、陳宥呈、廖泓閔')

    old_data = pd.DataFrame(old_data)

    train = reportData[0]
    test = reportData[1]
    x_train = reportData[2]
    x_test = reportData[3]
    y_train = reportData[4]
    y_test = reportData[5]
    module = reportData[6]

    dx.add_heading('訓練模型 - ' + choose, 1)

    dx.add_paragraph('\n選擇要訓練的模型：' + choose)
    dx.add_paragraph('要訓練的參數：')
    doc_table(dx, pd.DataFrame(train_list, columns = {'info'}))
    dx.add_paragraph('')
    dx.add_paragraph('要訓練的結果：' + choose)

    dx.add_heading('一、資料前處理')
    dx.add_paragraph('\n1. 讀取原本的資料 : ' + str(old_data.shape) + " [前15項]")
    doc_table(dx, old_data.head(15))
    dx.add_paragraph('')

    x_data = old_data[train_list]
    y_data = old_data[predict]

    train_data = pd.concat([x_data, y_data], axis=1)

    dx.add_paragraph('2. 將資料只留下要訓練的欄位 : ' + str(train_data.shape) + " [前15項]")
    doc_table(dx, train_data.head(15))
    dx.add_paragraph('')

    replace(train_data, train_list)

    dx.add_paragraph('3. 轉換訓練資料的欄位 : ' + str(train_data.shape) + " [前15項]")
    doc_table(dx, train_data.head(15))
    dx.add_paragraph('')

    dx.add_paragraph('4. 切割資料為訓練集與測試集，以 7 : 3 做切割點')
    dx.add_paragraph('')

    dx.add_paragraph('訓練用資料：' + str(train.shape) + " [前15項]")
    doc_table(dx, train.head(15))
    dx.add_paragraph('')
    dx.add_paragraph('測試用資料：' + str(test.shape) + " [前15項]")
    doc_table(dx, test.head(15))

    dx.add_heading('二、訓練模型')
    dx.add_paragraph('')

    y_pred = module.predict(x_test)


    if choose == "回歸模型" :

        variable_name = pd.Series(x_train.columns)
        variable = round(pd.Series(module.coef_), 3)
        variable_table = pd.concat([variable_name, variable], axis=1)
        variable_table.columns = ["name", "value"]

        dx.add_paragraph('\n參數訓練結果')
        doc_table(dx, variable_table)
        dx.add_paragraph('')

        regression = predict + " = "

        for i in range(0, len(variable)):

            regression = regression + variable_name[i]
            regression = regression + " x ( "
            regression = regression + str(round(variable[i], 3)) + " ) "

            if i != len(variable) - 1:
                regression = regression + " + "

        dx.add_paragraph('')
        dx.add_paragraph("回歸式：")
        dx.add_paragraph(regression)
        dx.add_paragraph('')

    elif choose == "NN模型":

        dx.add_paragraph('\n神經網路架構圖')
        dx.add_picture('../img/' + predict + 'Model_plot.png')
        dx.add_paragraph('')
        dx.add_paragraph("神經網路訓練 MSE 誤差")
        dx.add_picture('../nn_mse_error.jpg')

        y_test = y_test.reset_index(drop=True)
        y_pred = pd.DataFrame(y_pred, columns=['Pre'])['Pre']

    elif choose == "決策樹":

        dx.add_paragraph('\n決策樹架構圖')
        dx.add_picture('../img/tree_structure.png', width=Mm(277))
        dx.add_paragraph('圖片有點小，可以到網頁上下載原圖')
        dx.add_paragraph('')

    plt.plot([0, 20], [0, 20], color='red')
    plt.scatter(y_test, y_pred, color='blue')
    plt.title('Predict and Actual score')
    plt.xlabel('Actual score')
    plt.ylabel('Predict score')
    plt.savefig('distributed.jpg')

    dx.add_paragraph("預測成績分布")
    dx.add_picture(r'distributed.jpg')

    dx.add_paragraph('')
    dx.add_paragraph('模型誤差如下')

    dx.add_paragraph("mean_squared_error = {:.3}".format(mean_squared_error(y_pred, y_test)))
    dx.add_paragraph("mean_absolute_error = {:.3}".format(mean_absolute_error(y_pred, y_test)))
    dx.add_paragraph("median_absolute_error = {:.3}".format(median_absolute_error(y_pred, y_test)))

    dx.add_paragraph('')

    dx.add_paragraph('使用 r square 計算誤差')
    dx.add_paragraph('r square = {:.3}'.format(r2_score(y_pred, y_test)))

    dx.add_paragraph('')

    c = abs(y_pred - y_test)

    dx.add_paragraph("\n實際誤差分布百分比")
    dx.add_paragraph("error <= 0.5,\t共 " + str(len(c[c <= 0.5])) + " 筆資料\t\t約 {:.2%}".format(len(c[c <= 0.5]) / len(c)))
    dx.add_paragraph("error <= 1.5,\t共 " + str(len(c[c <= 1.5])) + " 筆資料\t\t約 {:.2%}".format(len(c[c <= 1.5]) / len(c)))
    dx.add_paragraph("error <= 2.5,\t共 " + str(len(c[c <= 2.5])) + " 筆資料\t\t約 {:.2%}".format(len(c[c <= 2.5]) / len(c)))

    dx.add_heading("分析報告結束")

    dx.save('report.docx')
